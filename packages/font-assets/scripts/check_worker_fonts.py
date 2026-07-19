#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any


EXPECTED_FONT_COUNT = 15


def load_manifest(path: Path) -> list[dict[str, Any]]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    fonts = payload.get("fonts")
    if not isinstance(fonts, list) or len(fonts) != EXPECTED_FONT_COUNT:
        raise RuntimeError(f"expected {EXPECTED_FONT_COUNT} font families")
    return fonts


def check_fontconfig(fonts: list[dict[str, Any]], font_root: Path) -> None:
    root = font_root.resolve()
    for font in fonts:
        family = str(font["family"])
        matched = subprocess.check_output(
            ["fc-match", "-f", "%{file}\n", family], text=True
        ).splitlines()[0]
        resolved = Path(matched).resolve()
        if not resolved.is_relative_to(root):
            raise RuntimeError(f"{family} resolved outside Orbit assets: {resolved}")
        print(f"fontconfig: {family} -> {resolved.name}")


def normalize_font_name(value: str) -> str:
    return re.sub(r"[^a-z0-9]", "", value.casefold())


def libreoffice_smoke(fonts: list[dict[str, Any]]) -> None:
    import fitz  # type: ignore[import-untyped]
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.oxml.ns import qn

    with tempfile.TemporaryDirectory(prefix="orbit-font-smoke-") as temp_dir:
        output_dir = Path(temp_dir)
        input_path = output_dir / "orbit-fonts.docx"
        document = Document()
        for index, font in enumerate(fonts):
            if index:
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            family = str(font["family"])
            text = "한글 발표 폰트 Orbit 123" if font["supportsKorean"] else "Orbit presentation 123"
            run = document.add_paragraph().add_run(text)
            run.font.name = family
            font_properties = run._element.get_or_add_rPr().rFonts
            if font_properties is None:
                raise RuntimeError(f"could not set the Word font for {family}")
            font_properties.set(qn("w:eastAsia"), family)
        document.save(str(input_path))

        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(input_path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            timeout=120,
        )
        pdf_path = output_dir / "orbit-fonts.pdf"
        if not pdf_path.is_file():
            raise RuntimeError("LibreOffice did not produce the font smoke PDF")

        pdf = fitz.open(pdf_path)
        if len(pdf) != len(fonts):
            raise RuntimeError("LibreOffice font smoke page count is inconsistent")
        for page, font in zip(pdf, fonts, strict=True):
            family = str(font["family"])
            expected = normalize_font_name(family)
            embedded = {
                normalize_font_name(str(item[3])) for item in page.get_fonts(full=True)
            }
            if not any(
                expected in candidate or candidate in expected for candidate in embedded
            ):
                raise RuntimeError(
                    f"LibreOffice substituted {family}: {sorted(embedded)}"
                )
            print(f"libreoffice: {family} -> {sorted(embedded)}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", type=Path, required=True)
    parser.add_argument("--font-root", type=Path, required=True)
    parser.add_argument("--check-fontconfig", action="store_true")
    parser.add_argument("--libreoffice-smoke", action="store_true")
    args = parser.parse_args()
    fonts = load_manifest(args.manifest)
    if args.check_fontconfig:
        check_fontconfig(fonts, args.font_root)
    if args.libreoffice_smoke:
        libreoffice_smoke(fonts)


if __name__ == "__main__":
    main()
